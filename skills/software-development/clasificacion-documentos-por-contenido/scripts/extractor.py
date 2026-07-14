#!/usr/bin/env python3
"""
extractor.py - Extraccion estructurada de resoluciones judiciales con Qwen 2.5 local.
Requiere: llama-server corriendo con Qwen 2.5 7B Q4_K_M.

Uso:
  python3 extractor.py --test              # Prueba con 5 samples
  python3 extractor.py --file ruta.pdf     # Documento individual
  python3 extractor.py --batch dir/        # Lote completo
  python3 extractor.py --no-server         # Usar server ya corriendo
"""

import os, sys, json, time, glob, subprocess, re, argparse
import csv
from concurrent.futures import ProcessPoolExecutor, as_completed
from pathlib import Path

BASE_DIR = Path(__file__).parent
MODEL_PATH = BASE_DIR / "modelos" / "qwen-7b-q4_k_m.gguf"
LLAMA_SERVER = BASE_DIR / "llama.cpp" / "build" / "bin" / "llama-server"
SAMPLES_DIR = BASE_DIR / "samples"
OUTPUT_DIR = BASE_DIR / "output"

# Atento: Quadro T1000 (4GB VRAM) OOM con 24 layers.
# 20 layers deja ~700 MB de headroom.
N_GPU_LAYERS = 20
N_THREADS = 12
N_CTX = 8192
BATCH_SIZE = 512

SYSTEM_PROMPT = """Eres un asistente legal experto en indexacion de jurisprudencia peruana.
Analiza la resolucion judicial proporcionada y extrae la informacion en formato JSON.
Debes devolver UNICAMENTE un objeto JSON valido, sin texto adicional."""

# NOTA: Las dobles llaves {{ }} son ESCAPES de .format() → producen { } literales en el JSON.
# {texto} (single brace) es el placeholder real para el texto.
USER_PROMPT_TPL = """Analiza esta resolucion judicial y genera un JSON con la siguiente estructura exacta:
{{
  "resumen_hechos": "Sintesis objetiva de los hechos relevantes (1-2 parrafos)",
  "resumen_problema": "Problema juridico central a resolver",
  "resumen_fallo": "Decision final del tribunal",
  "entidades_clave": {{
     "jueces_magistrados": ["Nombre del juez"],
     "demandantes_accionantes": ["Nombre o empresa"],
     "demandados_accionados": ["Nombre o empresa"],
     "leyes_y_articulos_citados": ["Ley X, Art. Y"],
     "conceptos_legales_clave": ["Concepto legal"]
  }}
}}

Instrucciones:
- Si una entidad no existe, usa arreglo vacio []
- Normaliza nombres de leyes
- Vocabulario tecnico-legal peruano
- Responde SOLO con el JSON, sin explicaciones

TEXTO DE LA RESOLUCION:
{texto}"""


def extraer_texto(path):
    path_str = str(path)
    lower = path_str.lower()
    try:
        if lower.endswith('.pdf'):
            res = subprocess.run(['pdftotext', '-layout', path_str, '-'],
                               capture_output=True, text=True, timeout=30)
            return res.stdout[:8000]
        elif lower.endswith('.docx'):
            from docx import Document
            doc = Document(path_str)
            return ' '.join(p.text for p in doc.paragraphs)[:8000]
        elif lower.endswith('.doc'):
            try:
                from docx import Document
                doc = Document(path_str)
                return ' '.join(p.text for p in doc.paragraphs)[:8000]
            except:
                # .doc legacy: probar catdoc, antiword, fallback a texto crudo
                try:
                    res = subprocess.run(['catdoc', path_str], capture_output=True, text=True, timeout=10)
                    if res.stdout:
                        return res.stdout[:8000]
                except FileNotFoundError:
                    pass
                try:
                    res = subprocess.run(['antiword', path_str], capture_output=True, text=True, timeout=10)
                    if res.stdout:
                        return res.stdout[:8000]
                except FileNotFoundError:
                    pass
                return ""
    except Exception as e:
        return f"[ERROR: {e}]"


def query_qwen(texto, server_url="http://127.0.0.1:8080"):
    prompt = USER_PROMPT_TPL.format(texto=texto[:7000])
    try:
        import requests
        resp = requests.post(f"{server_url}/v1/chat/completions", json={
            "model": "qwen",
            "messages": [
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.1,
            "max_tokens": 1024,
        }, timeout=120)
        content = resp.json()["choices"][0]["message"]["content"]
        content = content.strip()
        if content.startswith("```"):
            content = content.split("```")[1]
            content = re.sub(r'^json\s*', '', content, flags=re.IGNORECASE)
        return json.loads(content.strip())
    except Exception as e:
        return {"error": str(e)}


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="QwenLegalExtractor")
    parser.add_argument("--test", action="store_true")
    parser.add_argument("--file", type=str)
    parser.add_argument("--batch", type=str)
    parser.add_argument("--no-server", action="store_true")
    args = parser.parse_args()
    
    if args.test:
        print("Modo prueba: revisa samples/")
    elif args.file:
        print(procesar_documento(args.file))
    elif args.batch:
        print(f"Procesando {args.batch}")
